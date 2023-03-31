from flask import Flask, render_template, request, redirect, url_for, session, flash
from flask_wtf import FlaskForm
from flask_wtf.file import FileField, FileRequired, FileAllowed
from wtforms import StringField, SubmitField
from wtforms.validators import DataRequired, Email
from flask_uploads import UploadSet, configure_uploads
import pandas as pd
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Inches
from flask_login import login_required
import os, requests, shutil, json, googlemaps, smtplib, ssl
from datetime import datetime
from PIL import Image
from IPython.display import HTML
from email.message import EmailMessage


app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret_key'
app.config['UPLOADED_CSV_DEST'] = 'uploads/csv'
#   DATABASE CONFIGURATION
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///userlogs.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
#   FILE UPLOAD CONFIGURATION
csv_uploads = UploadSet('csv')
configure_uploads(app, csv_uploads)

class User(db.Model):
    __tablename__ = 'userlogs'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20))
    email = db.Column(db.String(100))
    csv_file = db.Column(db.String(100), nullable=False)
    added = db.Column(db.DateTime, default=datetime.now)

    def __init__(self, username, email, csv_file):
        self.username = username
        self.email = email
        self.csv_file = csv_file

def init_db():
    db.drop_all()
    db.create_all()
    new_user = User(username='user1', email='test@test.com', csv_file='test.csv')
    db.session.add(new_user)    
    db.session.commit()

@app.route('/clear_db', methods=['POST'])
def clear_db():
    all_users = db.session.query(User).all()
    
    for user in all_users:
        db.session.delete(user)
    db.session.commit()
    return render_template('logs.html')

class MyForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired(), Email()])
    csv_file = FileField('CSV File', validators=[
        FileRequired(),
        FileAllowed(['csv'],'FILE FORMATE MUST BE .CSV')
    ])
    submit = SubmitField('Submit')

@app.route('/', methods=['GET', 'POST'])
def index():
    session['logged_in'] = False
    form = MyForm()
    if form.validate_on_submit():
        username = form.username.data
        email = form.email.data
        csv_filename = csv_uploads.save(form.csv_file.data)
        #GETTING THE ADRESSES
        new_user = User(username=form.username.data, email=form.email.data, csv_file=csv_filename)
        db.session.add(new_user)
        db.session.commit()   
        session['email'] = email
        formatted_address(csv_filename)
        session['logged_in'] = True
        return redirect(url_for('success'))
    return render_template('index.html', form=form)

@app.route('/formatted')
def formatted():
    return render_template('pano.html')

@app.route('/success')
def success():
    return render_template('success.html')

@app.route('/logs')
def logs():
    data=User.query.all()
    return render_template('logs.html', data=data)


@app.route('/send_email', methods=['POST'])
def send_email():
    email_sender = 'disformatter@gmail.com'
    email_password = 'bqcwnzqvxhcryylq'
    email_reciever = session.get('email')

    subject = 'Document of Formatted Addresses'
    body = 'Please find the formatted addresses document attached.'

    em = EmailMessage()

    em['From'] = email_sender
    em['To'] = email_reciever
    em['Subject'] = subject
    em.set_content(body)

    with open('output.docx', 'rb') as f:
        file_data = f.read()
    em.add_attachment(file_data, maintype='application', subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', filename='output.docx')
    context = ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login(email_sender, email_password)
        smtp.sendmail(email_sender, email_reciever, em.as_string())
    return render_template('sent.html')

@app.before_first_request
def empty_folders():
    upload_folder = app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads', 'csv')
    temp_folder = app.config['TEMP_FOLDER'] = os.path.join(os.getcwd(), 'static')
    for folder in [upload_folder, temp_folder]:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))


def formatted_address(csv_filename):
    locations = pd.read_csv('uploads/csv/' + csv_filename)
    gmaps = googlemaps.Client(key='AIzaSyBwP_5ZGFGEhgo1Zc9cxW5l2jjEz5-gd1o')

    approx_address = []

    for i in range(0,len(locations)):
        loc1 = locations[['latitude (deg)', 'longitude (deg)']].iloc[i]
        reverse_geocode_result = gmaps.reverse_geocode((loc1['latitude (deg)'], loc1['longitude (deg)']))
        jsonString = json.dumps(reverse_geocode_result)
        jsonFile = open("data.json", "w")
        jsonFile.write(jsonString)
        jsonFile.close()
        with open('data.json', 'r', encoding='utf-8') as f:
            my_data = json.load(f)
            f_address = my_data[0]['formatted_address']
            approx_address.append(f_address)

    locations['approx-address'] = approx_address
    print(csv_filename)
    locations.to_csv('outputs/'+ csv_filename, index=False)
    pano(csv_filename)

def pano(csv_filename):
    locations_pano = pd.read_csv('outputs/' + csv_filename)
    headings = [0, 90, 180, 270]
    locations_pano['img_source'] = locations_pano['approx-address']
    for i in range(0, len(locations_pano)):
        lat1 = locations_pano.iloc[i]['latitude (deg)']
        long1 = locations_pano.iloc[i]['longitude (deg)']
        address = locations_pano.iloc[i]['approx-address']
        for heading in headings:
            input_heading = heading
            url = "https://maps.googleapis.com/maps/api/streetview?location={},{}&size=640x640&pitch=0&fov=90&heading={}&key=AIzaSyBwP_5ZGFGEhgo1Zc9cxW5l2jjEz5-gd1o".format(lat1, long1, input_heading)
            response = requests.get(url,stream=True)
            if response.status_code == 200:
                with open(f"standalone_images/{heading}.jpg", "wb") as f:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, f)

    ## CREATING PANO IMAGE BY STITCHING
        filenames = ["standalone_images/0.jpg", "standalone_images/90.jpg", "standalone_images/180.jpg", "standalone_images/270.jpg"]

        images = [Image.open(filename) for filename in filenames]
        widths, heights = zip(*(i.size for i in images))

        total_width = sum(widths)
        max_height = max(heights)

        pano_image = Image.new("RGB", (total_width, max_height), color="white")

        x_offset = 0
        for x, image in enumerate(images):
            pano_image.paste(image, (x_offset, 0))
            x_offset += widths[x]
        ## OUTPUTTING PANO IMAGES WITH ADDRESS NAME AS FILE NAME
        filename = str(i) +'.jpg'
        pano_image.save('static/' + filename)
        width = 1280
        height = 320
        locations_pano['img_source'].iloc[i] = r"<img src='{{ url_for('static', filename=" + repr(filename) + ") }}'>"
    to_word(csv_filename)
    # TO HTML
    locations_html = locations_pano.drop(columns=['GPS week', 'GPS second', 'solution status', 'height (m)'])
    result_html = locations_html.to_html()
    result_html_replaced = result_html.replace('&lt;','<').replace('&gt;','>')
    text_file = open('templates/pano.html', "w")
    text_file.write('{% extends "base.html" %}')
    text_file.write('{% block content %}')
    text_file.write(result_html_replaced)
    text_file.write('{% endblock %}')
    text_file.close()

def to_word(csv_filename):
    locations_toword = pd.read_csv('outputs/' + csv_filename)
    locations_toword = locations_toword.drop(columns=['GPS week', 'GPS second', 'solution status', 'height (m)'])
    doc = Document()
    table = doc.add_table(rows=1, cols=len(locations_toword.columns))

    header = table.rows[0].cells
    for i in range(len(locations_toword.columns)):
        header[i].text = locations_toword.columns[i]
    header[-1].text = 'Image'

    for index, row in locations_toword.iterrows():
        row_cells = table.add_row().cells
        for i in range(len(locations_toword.columns)):
            row_cells[i].text = str(row[i])
        image_cell = row_cells[-1]
        image_cell.add_paragraph().add_run().add_picture('static/' + str(index) + '.jpg', width=Inches(4.0), height=Inches(1.8))

    doc.save('output.docx')

if __name__ == '__main__':
    app.run(debug=True)


